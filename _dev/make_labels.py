from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_label_doc(filename, header_text, fields):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.0)
        section.bottom_margin = Cm(1.0)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            cell.width = Cm(9.0)
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            cell.paragraphs[0].clear()
            p_header = cell.paragraphs[0]
            p_header.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_left = p_header.add_run(header_text)
            run_left.bold = True
            run_left.font.size = Pt(11)
            run_left.font.name = 'Arial'
            run_right = p_header.add_run("    LabTrack")
            run_right.font.size = Pt(7)
            run_right.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            run_right.font.name = 'Arial'
            pPr = p_header._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single', qn('w:sz'): '6',
                qn('w:space'): '1', qn('w:color'): '333333'
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            for field in fields:
                p = cell.add_paragraph()
                p.space_before = Pt(1)
                p.space_after = Pt(1)
                run = p.add_run(field + ":")
                run.bold = False
                run.font.size = Pt(9)
                run.font.name = 'Arial'
    doc.save(filename)
    print(f"Criado: {filename}")

create_label_doc("C:\\labtrack\\ProjetoNovoControleLab\\defeito_mesa.docx",
    "DEF. MESA", ["Protocolo", "Cliente", "Pedido", "Data Entrada", "Vendedor", "Defeito"])

create_label_doc("C:\\labtrack\\ProjetoNovoControleLab\\defeito_equipamento.docx",
    "DEF. EQUIP.", ["Protocolo", "Tipo Protocolo", "N° Serie", "Data"])
